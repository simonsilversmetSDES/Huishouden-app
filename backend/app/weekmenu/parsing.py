"""De drie recept-parsers achter POST /api/weekmenu/recipes/parse (Fase 2).

1. URL met schema.org-metadata via ``recipe-scrapers`` (geen AI nodig);
2. URL-fallback: HTML strippen tot tekst en server-side naar de Claude API;
3. afbeelding/screenshot (base64) naar de Claude API (vision).

Alle drie geven hetzelfde bewerkbare ``ParsedRecipe`` terug; er wordt hier NOOIT
opgeslagen. De Anthropic-key komt uit backend/.env en blijft server-side.
"""

import base64
import io
import json
import re

import anthropic
import docx
from bs4 import BeautifulSoup
from pypdf import PdfReader
from pypdf.errors import PdfReadError
from recipe_scrapers import scrape_html

from app.config import Settings
from app.weekmenu.errors import WeekmenuError
from app.weekmenu.schemas import ParsedIngredient, ParsedRecipe
from app.weekmenu.url_security import MAX_HTML_BYTES, fetch_url

MAX_CLAUDE_TEXT_CHARS = 50_000
CLAUDE_MAX_TOKENS = 3000

# JSON-contract uit WEEKMENU_BUILD.md (Engelse sleutels, zelfde structuur) + servings.
_JSON_FORMAT = """{
  "title": "...",
  "description": "...",
  "photo_url": null,
  "source_url": null,
  "servings": null,
  "ingredients": [
    {"name": "...", "quantity": "... of null", "unit": "... of null"}
  ]
}"""

_SERVINGS_INSTRUCTION = (
    'Bepaal ook voor hoeveel personen het recept bedoeld is en zet dat als geheel getal '
    'in "servings"; gebruik null als dit nergens vermeld staat.'
)

_TEXT_PROMPT = (
    "Hieronder staat de tekst van een webpagina met een recept. Extraheer het recept "
    "en antwoord met UITSLUITEND geldige JSON in exact dit formaat, zonder uitleg en "
    f"zonder markdown-fences:\n{_JSON_FORMAT}\n"
    "Zet de bereidingsstappen als doorlopende tekst (één stap per regel) in "
    '"description". Hou hoeveelheden en eenheden apart. '
    f"{_SERVINGS_INSTRUCTION} Gebruik null voor andere onbekende "
    "velden. Antwoord in het Nederlands.\n\nPAGINATEKST:\n"
)

_IMAGE_PROMPT = (
    "Op deze afbeelding staat een recept (foto of screenshot). Extraheer het recept "
    "en antwoord met UITSLUITEND geldige JSON in exact dit formaat, zonder uitleg en "
    f"zonder markdown-fences:\n{_JSON_FORMAT}\n"
    "Zet de bereidingsstappen als doorlopende tekst (één stap per regel) in "
    '"description". Hou hoeveelheden en eenheden apart. '
    f"{_SERVINGS_INSTRUCTION} Gebruik null voor andere onbekende "
    "velden. Antwoord in het Nederlands."
)

# Heuristiek om "500 g bloem" te splitsen; eenheden die we als eenheid herkennen.
_KNOWN_UNITS = {
    "g",
    "gr",
    "gram",
    "kg",
    "kilo",
    "kilogram",
    "mg",
    "ml",
    "cl",
    "dl",
    "l",
    "liter",
    "el",
    "eetlepel",
    "eetlepels",
    "tl",
    "theelepel",
    "theelepels",
    "snufje",
    "snufjes",
    "mespunt",
    "mespuntje",
    "teentje",
    "teentjes",
    "stuk",
    "stuks",
    "stukje",
    "stukjes",
    "blik",
    "blikje",
    "blikjes",
    "zakje",
    "zakjes",
    "pak",
    "pakje",
    "pakjes",
    "bosje",
    "bosjes",
    "takje",
    "takjes",
    "plak",
    "plakken",
    "plakjes",
    "scheut",
    "scheutje",
    "kop",
    "kopje",
    "kopjes",
    "cup",
    "cups",
    "tbsp",
    "tsp",
    "oz",
    "lb",
}
_INGREDIENT_RE = re.compile(
    r"^\s*(?P<qty>\d+(?:[.,]\d+)?(?:\s*[-–/]\s*\d+(?:[.,]\d+)?)?|½|¼|¾)\s*"
    r"(?:(?P<unit>[A-Za-zéë]+\.?)\s+)?(?P<rest>.+)$"
)

# recipe-scrapers geeft yields() als vrije tekst terug ("4 servings", "4 porties", "4-6 personen").
_YIELDS_RE = re.compile(r"(\d+)")

# Sommige sites (bv. VRT Dagelijkse Kost) zetten in hun schema.org-JSON-LD enkel de
# eerste paar stappen; het volledige stappenplan zit als dehydrated React Query-data
# in een Next.js flight-payload (``self.__next_f.push(...)``) verderop in de pagina.
_NEXT_FLIGHT_RE = re.compile(r"self\.__next_f\.push\(\[1,\s*(\".*?\")\]\)", re.S)
_HEX_CHUNK_PREFIX_RE = re.compile(r"^[0-9a-fA-F]+:")


def _find_recipe_parts(obj: object) -> dict | None:
    """Zoekt recursief naar de dict met ``recipeParts`` in een Next.js flight-payload."""
    if isinstance(obj, dict):
        if "recipeParts" in obj and "ingredients" in obj:
            return obj
        for value in obj.values():
            found = _find_recipe_parts(value)
            if found is not None:
                return found
    elif isinstance(obj, list):
        for item in obj:
            found = _find_recipe_parts(item)
            if found is not None:
                return found
    return None


def _extract_nextjs_flight_description(html: str) -> str | None:
    """Best-effort: volledige bereidingswijze uit een Next.js flight-payload.

    Geeft None terug bij twijfel — de schema.org-tekst blijft dan gewoon staan.
    """
    for match in _NEXT_FLIGHT_RE.finditer(html):
        try:
            chunk = json.loads(match.group(1))
        except json.JSONDecodeError:
            continue
        body = _HEX_CHUNK_PREFIX_RE.sub("", chunk, count=1)
        try:
            payload = json.loads(body)
        except json.JSONDecodeError:
            continue
        recipe = _find_recipe_parts(payload)
        if recipe is None:
            continue
        lines: list[str] = []
        for part in recipe.get("recipeParts") or []:
            if part.get("title"):
                lines.append(f"{part['title']}:")
            for instruction in part.get("instructions") or []:
                description = instruction.get("description")
                if description:
                    lines.append(description)
        if lines:
            return "\n".join(lines)
    return None


def split_ingredient_line(line: str) -> ParsedIngredient:
    """Splits een ingrediëntregel in hoeveelheid/eenheid/naam; lukt dat niet → alles in name."""
    line = " ".join(line.split())
    match = _INGREDIENT_RE.match(line)
    if not match:
        return ParsedIngredient(name=line)
    qty = match.group("qty")
    unit = match.group("unit")
    rest = match.group("rest")
    if unit and unit.rstrip(".").lower() not in _KNOWN_UNITS:
        # Geen echte eenheid ("2 uien") — het woord hoort bij de naam.
        rest = f"{unit} {rest}"
        unit = None
    return ParsedIngredient(name=rest.strip(), quantity=qty, unit=unit)


def _try_schema_org(html: str, url: str) -> ParsedRecipe | None:
    """Parser 1: recipe-scrapers op schema.org/JSON-LD; None bij onbruikbaar resultaat."""
    try:
        scraper = scrape_html(html, org_url=url, supported_only=False)
    except Exception:
        return None

    def _get(getter: str) -> object | None:
        try:
            return getattr(scraper, getter)()
        except Exception:
            return None

    title = _get("title")
    raw_ingredients = _get("ingredients")
    if not title or not raw_ingredients:
        return None
    instructions = _get("instructions") or ""
    flight_instructions = _extract_nextjs_flight_description(html)
    if flight_instructions and len(flight_instructions) > len(instructions):
        instructions = flight_instructions
    image = _get("image")
    yields = _get("yields")
    servings_match = _YIELDS_RE.search(str(yields)) if yields else None
    return ParsedRecipe(
        title=str(title).strip(),
        description=str(instructions).strip(),
        photo_url=str(image) if image else None,
        source_url=url,
        servings=int(servings_match.group(1)) if servings_match else None,
        ingredients=[
            split_ingredient_line(str(line)) for line in raw_ingredients if str(line).strip()
        ],
    )


def _html_to_text(html: str) -> str:
    """Strip HTML tot platte tekst voor de Claude-fallback (scripts/nav eruit, gecapt)."""
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "svg", "nav", "footer", "header", "form"]):
        tag.decompose()
    text = soup.get_text(separator="\n").replace("\xa0", " ")  # nbsp normaliseren
    lines = [" ".join(line.split()) for line in text.splitlines()]
    text = "\n".join(line for line in lines if line)
    return text[:MAX_CLAUDE_TEXT_CHARS]


def extract_docx_text(data: bytes) -> str:
    """Platte tekst uit een .docx: paragrafen + tabelcellen (recepten staan soms in tabellen)."""
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise WeekmenuError(
            400, "invalid_document", "Het Word-bestand kon niet gelezen worden."
        ) from exc
    lines = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())
    return "\n".join(lines)[:MAX_CLAUDE_TEXT_CHARS]


def extract_pdf_text(data: bytes) -> str:
    """Platte tekst uit een PDF, pagina per pagina (geen OCR — vereist een tekstlaag)."""
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except PdfReadError as exc:
        raise WeekmenuError(400, "invalid_document", "De PDF kon niet gelezen worden.") from exc
    return "\n".join(pages)[:MAX_CLAUDE_TEXT_CHARS]


def get_claude_client(settings: Settings) -> anthropic.Anthropic:
    """Gedeeld met ingredient_categorization.py — zelfde 503 als geen key geconfigureerd is."""
    if not settings.anthropic_api_key:
        raise WeekmenuError(
            503,
            "ai_unavailable",
            "AI-parsing is niet geconfigureerd (geen ANTHROPIC_API_KEY in backend/.env).",
        )
    return anthropic.Anthropic(api_key=settings.anthropic_api_key)


def _claude_call(settings: Settings, content: list[dict]) -> ParsedRecipe:
    client = get_claude_client(settings)
    try:
        response = client.messages.create(
            model=settings.anthropic_model,
            max_tokens=CLAUDE_MAX_TOKENS,
            messages=[{"role": "user", "content": content}],
        )
    except anthropic.AnthropicError as exc:
        raise WeekmenuError(
            502, "ai_failed", "Het recept kon niet automatisch gelezen worden."
        ) from exc
    raw = "".join(block.text for block in response.content if block.type == "text")
    return _parse_claude_json(raw)


def _parse_claude_json(raw: str) -> ParsedRecipe:
    text = raw.strip()
    # Strip eventuele ```json-fences ondanks de prompt-instructie.
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\n?", "", text)
        text = re.sub(r"\n?```$", "", text.strip())
    try:
        data = json.loads(text)
        return ParsedRecipe.model_validate(data)
    except (json.JSONDecodeError, ValueError) as exc:
        raise WeekmenuError(
            502,
            "ai_invalid_response",
            "Het recept kon niet automatisch gelezen worden; probeer het opnieuw.",
        ) from exc


def call_claude_text(text: str, settings: Settings) -> ParsedRecipe:
    """Parser 2: gestripte paginatekst naar Claude, strikt JSON terug."""
    return _claude_call(settings, [{"type": "text", "text": _TEXT_PROMPT + text}])


def call_claude_image(image_base64: str, media_type: str, settings: Settings) -> ParsedRecipe:
    """Parser 3: afbeelding (base64) naar Claude vision, strikt JSON terug."""
    return _claude_call(
        settings,
        [
            {
                "type": "image",
                "source": {"type": "base64", "media_type": media_type, "data": image_base64},
            },
            {"type": "text", "text": _IMAGE_PROMPT},
        ],
    )


def parse_url(url: str, settings: Settings) -> ParsedRecipe:
    """URL-flow: HTML éénmaal ophalen → parser 1 (schema.org) → parser 2 (Claude)."""
    result = fetch_url(url, MAX_HTML_BYTES)
    html = result.content.decode("utf-8", errors="replace")
    parsed = _try_schema_org(html, result.final_url)
    if parsed is not None:
        return parsed
    recipe = call_claude_text(_html_to_text(html), settings)
    recipe.source_url = result.final_url
    return recipe


def parse_image(image_base64: str, media_type: str, settings: Settings) -> ParsedRecipe:
    """Afbeelding-flow: rechtstreeks naar Claude vision."""
    return call_claude_image(image_base64, media_type, settings)


_DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def parse_document(document_base64: str, media_type: str, settings: Settings) -> ParsedRecipe:
    """Document-flow (Word/PDF): tekst extraheren → zelfde Claude-tekstpad als de URL-fallback."""
    data = base64.b64decode(document_base64)
    text = extract_docx_text(data) if media_type == _DOCX_MEDIA_TYPE else extract_pdf_text(data)
    if not text.strip():
        raise WeekmenuError(
            422,
            "empty_document",
            "Kon geen tekst uit dit bestand halen (bv. een gescande PDF zonder tekstlaag).",
        )
    return call_claude_text(text, settings)
